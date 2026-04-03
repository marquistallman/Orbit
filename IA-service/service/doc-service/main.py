import os
from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document

app = FastAPI()


# Keep generated files in a dedicated directory mounted by docker-compose.
OUTPUT_DIR = os.getenv("OUTPUT_DIR", os.path.join(os.path.dirname(os.path.abspath(__file__)), "output"))

# crear carpeta si no existe
os.makedirs(OUTPUT_DIR, exist_ok=True)


class DocRequest(BaseModel):
    title: str
    content: str
    format: str = "docx"  # docx | pdf | both


class DocChangeRequest(BaseModel):
    file_name: str
    replacements: dict[str, str] = {}
    append_text: str = ""


@app.get("/")
def health():
    return {"status": "doc service running", "output_dir": OUTPUT_DIR}


@app.get("/files/{file_name}")
def download_file(file_name: str):
    file_path = os.path.join(OUTPUT_DIR, file_name)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(path=file_path, filename=file_name)


def _safe_filename(name: str) -> str:
    cleaned = "".join(ch for ch in name if ch.isalnum() or ch in (" ", "-", "_"))
    return cleaned.strip().replace(" ", "_") or "document"


def _write_docx(title: str, content: str, base_filename: str) -> str:
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    file_path = os.path.join(OUTPUT_DIR, f"{base_filename}.docx")
    doc.save(file_path)
    return file_path


def _write_pdf(title: str, content: str, base_filename: str) -> str:
    LETTER = __import__("reportlab.lib.pagesizes", fromlist=["LETTER"]).LETTER
    canvas = __import__("reportlab.pdfgen.canvas", fromlist=["Canvas"])
    file_path = os.path.join(OUTPUT_DIR, f"{base_filename}.pdf")
    pdf = canvas.Canvas(file_path, pagesize=LETTER)
    width, height = LETTER
    y = height - 72

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(72, y, title)
    y -= 28

    pdf.setFont("Helvetica", 11)
    for raw_line in content.splitlines() or [""]:
        line = raw_line
        while len(line) > 105:
            pdf.drawString(72, y, line[:105])
            line = line[105:]
            y -= 16
            if y < 72:
                pdf.showPage()
                pdf.setFont("Helvetica", 11)
                y = height - 72
        pdf.drawString(72, y, line)
        y -= 16
        if y < 72:
            pdf.showPage()
            pdf.setFont("Helvetica", 11)
            y = height - 72

    pdf.save()
    return file_path


def _must_exist(file_name: str) -> str:
    path = os.path.join(OUTPUT_DIR, file_name)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="File not found")
    return path


@app.post("/edit")
def create_doc(req: DocRequest):
    base_filename = _safe_filename(req.title)
    format_value = req.format.lower().strip()
    file_paths: list[str] = []

    if format_value in ("docx", "word", "both"):
        file_paths.append(_write_docx(req.title, req.content, base_filename))

    if format_value in ("pdf", "both"):
        file_paths.append(_write_pdf(req.title, req.content, base_filename))

    if not file_paths:
        return {
            "error": "Invalid format. Use docx, pdf, or both.",
            "accepted_formats": ["docx", "pdf", "both"]
        }

    return {
        "message": "Document created",
        "files": file_paths,
        "file_names": [os.path.basename(path) for path in file_paths],
        "format": format_value
    }


@app.post("/apply-changes")
def apply_changes(req: DocChangeRequest):
    original_path = _must_exist(req.file_name)
    _, ext = os.path.splitext(req.file_name.lower())

    if ext == ".docx":
        doc = Document(original_path)
        for para in doc.paragraphs:
            updated = para.text
            for old, new in req.replacements.items():
                updated = updated.replace(old, new)
            if updated != para.text:
                para.text = updated

        if req.append_text:
            doc.add_paragraph(req.append_text)

        out_name = f"{os.path.splitext(req.file_name)[0]}_edited.docx"
        out_path = os.path.join(OUTPUT_DIR, out_name)
        doc.save(out_path)
        return {
            "message": "Document edited",
            "source_file": req.file_name,
            "file_name": out_name,
            "file_path": out_path,
        }

    if ext == ".pdf":
        revised_text = req.append_text or "No textual changes provided"
        for old, new in req.replacements.items():
            revised_text += f"\nReplace: {old} -> {new}"

        out_name = f"{os.path.splitext(req.file_name)[0]}_edited.pdf"
        out_path = _write_pdf("Edited PDF", revised_text, os.path.splitext(out_name)[0])
        return {
            "message": "PDF revision generated",
            "source_file": req.file_name,
            "file_name": os.path.basename(out_path),
            "file_path": out_path,
            "note": "This creates a revised PDF output based on requested changes.",
        }

    raise HTTPException(status_code=400, detail="Unsupported file type")